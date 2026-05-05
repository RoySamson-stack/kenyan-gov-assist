"""
Extended Document Processors for various file formats
Supports: PDF, DOCX, XLSX, TXT, MD, EPUB, HTML
"""
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
import re

logger = logging.getLogger(__name__)


class BaseDocumentProcessor:
    """Base class for document processors."""
    
    def __init__(self):
        self.supported_extensions = []
    
    def can_process(self, file_path: str) -> bool:
        """Check if this processor can handle the file."""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        """Process file and return list of text chunks with metadata."""
        raise NotImplementedError


class TextDocumentProcessor(BaseDocumentProcessor):
    """Process plain text files (.txt, .md)."""
    
    def __init__(self):
        self.supported_extensions = ['.txt', '.md', '.rst', '.log']
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        chunks = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Split into chunks (by paragraphs)
            paragraphs = re.split(r'\n\s*\n', content)
            
            for i, para in enumerate(paragraphs):
                para = para.strip()
                if len(para) > 50:  # Skip very short paragraphs
                    chunks.append({
                        'content': para,
                        'metadata': {
                            'source': Path(file_path).name,
                            'chunk_index': i,
                            'file_type': 'text',
                            'paragraph_index': i
                        }
                    })
            
            logger.info(f"Processed text file: {file_path} -> {len(chunks)} chunks")
        
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
        
        return chunks


class ExcelDocumentProcessor(BaseDocumentProcessor):
    """Process Excel files (.xlsx, .xls)."""
    
    def __init__(self):
        self.supported_extensions = ['.xlsx', '.xls', '.csv']
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        chunks = []
        
        try:
            import pandas as pd
            
            if file_path.endswith('.csv'):
                dfs = [('Sheet1', pd.read_csv(file_path))]
            else:
                excel_file = pd.ExcelFile(file_path)
                dfs = [(sheet, excel_file.parse(sheet)) for sheet in excel_file.sheet_names]
            
            for sheet_name, df in dfs:
                # Convert dataframe to text representations
                chunks.append({
                    'content': f"Sheet: {sheet_name}\nColumns: {', '.join(df.columns.tolist())}",
                    'metadata': {
                        'source': Path(file_path).name,
                        'sheet': sheet_name,
                        'file_type': 'excel',
                        'chunk_type': 'header'
                    }
                })
                
                # Process rows in batches
                batch_size = 50
                for start in range(0, len(df), batch_size):
                    batch = df.iloc[start:start + batch_size]
                    
                    # Convert to text
                    text_rows = []
                    for idx, row in batch.iterrows():
                        row_text = ' | '.join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                        if row_text:
                            text_rows.append(f"Row {idx}: {row_text}")
                    
                    if text_rows:
                        chunks.append({
                            'content': '\n'.join(text_rows),
                            'metadata': {
                                'source': Path(file_path).name,
                                'sheet': sheet_name,
                                'file_type': 'excel',
                                'chunk_type': 'data',
                                'row_start': start,
                                'row_end': min(start + batch_size, len(df))
                            }
                        })
            
            logger.info(f"Processed Excel file: {file_path} -> {len(chunks)} chunks")
        
        except ImportError:
            logger.error("pandas not installed. Install with: pip install pandas openpyxl")
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {e}")
        
        return chunks


class WordDocumentProcessor(BaseDocumentProcessor):
    """Process Word documents (.docx)."""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.doc']
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        chunks = []
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract paragraphs
            current_chunk = []
            chunk_size = 500
            current_size = 0
            chunk_idx = 0
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                if current_size + len(text) > chunk_size and current_chunk:
                    chunks.append({
                        'content': ' '.join(current_chunk),
                        'metadata': {
                            'source': Path(file_path).name,
                            'chunk_index': chunk_idx,
                            'file_type': 'word',
                            'paragraph_style': para.style.name if hasattr(para, 'style') else 'normal'
                        }
                    })
                    chunk_idx += 1
                    current_chunk = []
                    current_size = 0
                
                current_chunk.append(text)
                current_size += len(text)
            
            # Last chunk
            if current_chunk:
                chunks.append({
                    'content': ' '.join(current_chunk),
                    'metadata': {
                        'source': Path(file_path).name,
                        'chunk_index': chunk_idx,
                        'file_type': 'word'
                    }
                })
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        table_text.append(row_text)
                
                if table_text:
                    chunks.append({
                        'content': '\n'.join(table_text),
                        'metadata': {
                            'source': Path(file_path).name,
                            'table_index': table_idx,
                            'file_type': 'word',
                            'chunk_type': 'table'
                        }
                    })
            
            logger.info(f"Processed Word file: {file_path} -> {len(chunks)} chunks")
        
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"Error processing Word file {file_path}: {e}")
        
        return chunks


class EPUBDocumentProcessor(BaseDocumentProcessor):
    """Process EPUB ebook files."""
    
    def __init__(self):
        self.supported_extensions = ['.epub']
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        chunks = []
        
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
            
            book = epub.read_epub(file_path)
            
            chunk_idx = 0
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    # Parse HTML content
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    
                    # Remove script/style tags
                    for tag in soup(['script', 'style']):
                        tag.decompose()
                    
                    text = soup.get_text(separator='\n')
                    text = re.sub(r'\n\s*\n', '\n\n', text)
                    
                    # Split into chunks
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        para = para.strip()
                        if len(para) > 100:
                            chunks.append({
                                'content': para,
                                'metadata': {
                                    'source': Path(file_path).name,
                                    'chunk_index': chunk_idx,
                                    'file_type': 'epub',
                                    'item_id': item.get_id(),
                                    'char_count': len(para)
                                }
                            })
                            chunk_idx += 1
            
            logger.info(f"Processed EPUB file: {file_path} -> {len(chunks)} chunks")
        
        except ImportError:
            logger.error("ebooklib/beautifulsoup4 not installed. Install with: pip install ebooklib beautifulsoup4")
        except Exception as e:
            logger.error(f"Error processing EPUB file {file_path}: {e}")
        
        return chunks


class HTMLDocumentProcessor(BaseDocumentProcessor):
    """Process HTML files."""
    
    def __init__(self):
        self.supported_extensions = ['.html', '.htm', '.xhtml']
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        chunks = []
        
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
            
            # Remove script/style
            for tag in soup(['script', 'style', 'nav', 'footer']):
                tag.decompose()
            
            # Get title
            title = soup.title.string if soup.title else Path(file_path).stem
            
            # Extract main content
            content = soup.get_text(separator='\n')
            content = re.sub(r'\n\s*\n', '\n\n', content)
            
            # Split into chunks
            paragraphs = content.split('\n\n')
            for i, para in enumerate(paragraphs):
                para = para.strip()
                if len(para) > 100:
                    chunks.append({
                        'content': para,
                        'metadata': {
                            'source': Path(file_path).name,
                            'title': title,
                            'chunk_index': i,
                            'file_type': 'html',
                            'char_count': len(para)
                        }
                    })
            
            logger.info(f"Processed HTML file: {file_path} -> {len(chunks)} chunks")
        
        except ImportError:
            logger.error("beautifulsoup4 not installed. Install with: pip install beautifulsoup4")
        except Exception as e:
            logger.error(f"Error processing HTML file {file_path}: {e}")
        
        return chunks


class DocumentProcessorFactory:
    """Factory to get appropriate processor for a file."""
    
    def __init__(self):
        self.processors = [
            TextDocumentProcessor(),
            ExcelDocumentProcessor(),
            WordDocumentProcessor(),
            EPUBDocumentProcessor(),
            HTMLDocumentProcessor(),
        ]
    
    def get_processor(self, file_path: str) -> Optional[BaseDocumentProcessor]:
        """Get the appropriate processor for a file."""
        for processor in self.processors:
            if processor.can_process(file_path):
                return processor
        return None
    
    def process_file(self, file_path: str) -> List[Dict[str, Any]]:
        """Process any supported file type."""
        processor = self.get_processor(file_path)
        if processor:
            return processor.process(file_path)
        else:
            logger.warning(f"No processor found for: {file_path}")
            return []
